from __future__ import annotations

import hashlib
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from fastapi import HTTPException, UploadFile, status
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models.news import Post
from app.models.post_documents import PostDocument


def _utc_now() -> datetime:
    return datetime.now(timezone.utc)


def _storage_root() -> Path:
    return Path(settings.onlyoffice_storage_dir).resolve()


def _uploads_root() -> Path:
    return Path(settings.upload_dir).resolve()


def _slug_or_fallback(post: Post) -> str:
    candidate = (post.slug or "").strip()
    return candidate or f"post-{post.id}"


def generate_document_key(post_id: int, version: int, changed_at: datetime | None = None) -> str:
    timestamp = changed_at or _utc_now()
    digest = hashlib.sha1(f"{post_id}:{version}:{timestamp.isoformat()}".encode("utf-8")).hexdigest()[:12]
    return f"post-{post_id}-v{version}-{digest}"


def resolve_absolute_file_path(file_path: str) -> Path:
    path = Path(file_path)
    if path.is_absolute():
        return path
    return Path.cwd() / path


def _build_file_name(post: Post, original_name: str | None = None) -> str:
    base_name = _slug_or_fallback(post)
    suffix = ".docx"
    if original_name:
        original_suffix = Path(original_name).suffix.lower()
        if original_suffix == ".docx":
            suffix = original_suffix
    return f"{base_name}{suffix}"


def _build_relative_file_path(post_id: int, file_name: str) -> str:
    return str(Path(settings.onlyoffice_storage_dir) / str(post_id) / file_name)


def _build_file_url(relative_file_path: str) -> str:
    public_base = (settings.onlyoffice_docx_public_base_url or "").strip().rstrip("/")
    absolute_path = resolve_absolute_file_path(relative_file_path)
    relative_to_storage = absolute_path.relative_to(_storage_root()).as_posix()
    if public_base:
        return f"{public_base}/{relative_to_storage}"

    callback_base = (settings.onlyoffice_callback_base_url or "").strip().rstrip("/")
    upload_prefix = settings.upload_url_prefix.rstrip("/")
    try:
        relative_to_uploads = absolute_path.relative_to(_uploads_root()).as_posix()
    except ValueError as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="ONLYOFFICE storage dir must live under upload_dir or ONLYOFFICE_DOCX_PUBLIC_BASE_URL must be configured.",
        ) from exc
    if not callback_base:
        return f"{upload_prefix}/{relative_to_uploads}"
    return f"{callback_base}{upload_prefix}/{relative_to_uploads}"


def _write_blank_document(post: Post, destination: Path) -> None:
    destination.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading(post.title or f"Post {post.id}", level=1)
    document.add_paragraph("Start writing the article content here.")
    document.save(destination)


def get_post_document(db: Session, post_id: int) -> PostDocument | None:
    return db.scalar(select(PostDocument).where(PostDocument.post_id == post_id))


def ensure_post_document(db: Session, post_id: int) -> PostDocument:
    post = db.get(Post, post_id)
    if not post:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Post not found.")

    existing = get_post_document(db=db, post_id=post_id)
    if existing and resolve_absolute_file_path(existing.file_path).exists():
        if not existing.file_url:
            existing.file_url = _build_file_url(existing.file_path)
            db.add(existing)
            db.commit()
            db.refresh(existing)
        return existing

    file_name = _build_file_name(post)
    relative_file_path = _build_relative_file_path(post_id=post.id, file_name=file_name)
    absolute_file_path = resolve_absolute_file_path(relative_file_path)
    _write_blank_document(post=post, destination=absolute_file_path)

    changed_at = _utc_now()
    document = existing or PostDocument(post_id=post.id, file_name=file_name, file_path=relative_file_path, version=1)
    document.file_name = file_name
    document.file_path = relative_file_path
    document.file_url = _build_file_url(relative_file_path)
    document.last_synced_at = changed_at
    document.document_key = generate_document_key(post_id=post.id, version=document.version or 1, changed_at=changed_at)

    db.add(document)
    db.commit()
    db.refresh(document)
    return document


def save_document_bytes(
    db: Session,
    document: PostDocument,
    file_bytes: bytes,
    file_name: str | None = None,
    increment_version: bool = True,
) -> PostDocument:
    post = db.get(Post, document.post_id)
    if not post:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Post not found.")

    next_file_name = file_name or document.file_name or _build_file_name(post)
    relative_file_path = _build_relative_file_path(post_id=post.id, file_name=next_file_name)
    absolute_file_path = resolve_absolute_file_path(relative_file_path)
    absolute_file_path.parent.mkdir(parents=True, exist_ok=True)
    absolute_file_path.write_bytes(file_bytes)

    changed_at = _utc_now()
    document.file_name = next_file_name
    document.file_path = relative_file_path
    document.file_url = _build_file_url(relative_file_path)
    document.last_synced_at = changed_at
    document.version = max(1, int(document.version or 1) + (1 if increment_version else 0))
    document.document_key = generate_document_key(post_id=post.id, version=document.version, changed_at=changed_at)

    db.add(document)
    db.commit()
    db.refresh(document)
    return document


async def replace_post_document_from_upload(db: Session, post_id: int, file: UploadFile) -> PostDocument:
    file_name = (file.filename or "").strip()
    if not file_name.lower().endswith(".docx"):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Only .docx files are supported.")

    document = ensure_post_document(db=db, post_id=post_id)
    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Uploaded file is empty.")

    return save_document_bytes(
        db=db,
        document=document,
        file_bytes=file_bytes,
        file_name=_build_file_name(db.get(Post, post_id), original_name=file_name),
        increment_version=True,
    )
